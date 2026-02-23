"""
Document Ingestion Service
Handles parsing and chunking of various document formats
"""

import os
from typing import List, Tuple
import PyPDF2
from docx import Document
import markdown
from pathlib import Path

class DocumentIngestor:
    """Handles document parsing and chunking"""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        """
        Initialize the document ingestor
        
        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse_pdf(self, file_path: str) -> str:
        """Parse PDF file and extract text"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
        except Exception as e:
            raise Exception(f"Error parsing PDF {file_path}: {str(e)}")
        return text
    
    def parse_docx(self, file_path: str) -> str:
        """Parse DOCX file and extract text"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX {file_path}: {str(e)}")
        return text
    
    def parse_txt(self, file_path: str) -> str:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error parsing TXT {file_path}: {str(e)}")
    
    def parse_markdown(self, file_path: str) -> str:
        """Parse Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_text = file.read()
                # Convert markdown to plain text
                html = markdown.markdown(md_text)
                # Simple HTML to text conversion
                import re
                text = re.sub('<[^<]+?>', '', html)
                return text
        except Exception as e:
            raise Exception(f"Error parsing Markdown {file_path}: {str(e)}")
    
    def parse_document(self, file_path: str) -> str:
        """
        Parse document based on file extension
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text content
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.parse_pdf(file_path)
        elif file_ext == '.docx':
            return self.parse_docx(file_path)
        elif file_ext == '.txt':
            return self.parse_txt(file_path)
        elif file_ext in ['.md', '.markdown']:
            return self.parse_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                if last_period > self.chunk_size // 2:
                    end = start + last_period + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
        
        return chunks
    
    def ingest_document(self, file_path: str) -> List[str]:
        """
        Ingest a document: parse and chunk
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of text chunks
        """
        text = self.parse_document(file_path)
        chunks = self.chunk_text(text)
        return chunks


# Global instance
ingestor = DocumentIngestor()
